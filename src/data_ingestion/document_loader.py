"""
Document Loader for RAG Pipeline

Supports loading documents from various formats including PDF, DOCX, TXT, and HTML.
Implements error handling and format validation.
"""

import os
import logging
from pathlib import Path
from typing import List, Dict, Any, Optional
from abc import ABC, abstractmethod

import PyPDF2
from docx import Document
from bs4 import BeautifulSoup
import requests

logger = logging.getLogger(__name__)


class BaseDocumentLoader(ABC):
    """Abstract base class for document loaders."""
    
    @abstractmethod
    def load(self, file_path: str) -> str:
        """Load document content from file path."""
        pass
    
    @abstractmethod
    def get_metadata(self, file_path: str) -> Dict[str, Any]:
        """Extract metadata from document."""
        pass


class PDFLoader(BaseDocumentLoader):
    """Loader for PDF documents."""
    
    def load(self, file_path: str) -> str:
        """Extract text content from PDF file."""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                text = ""
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
                return text.strip()
        except Exception as e:
            logger.error(f"Error loading PDF {file_path}: {e}")
            raise
    
    def get_metadata(self, file_path: str) -> Dict[str, Any]:
        """Extract PDF metadata."""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                return {
                    'num_pages': len(pdf_reader.pages),
                    'file_size': os.path.getsize(file_path),
                    'format': 'PDF'
                }
        except Exception as e:
            logger.error(f"Error extracting PDF metadata: {e}")
            return {}


class DOCXLoader(BaseDocumentLoader):
    """Loader for DOCX documents."""
    
    def load(self, file_path: str) -> str:
        """Extract text content from DOCX file."""
        try:
            doc = Document(file_path)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            return text.strip()
        except Exception as e:
            logger.error(f"Error loading DOCX {file_path}: {e}")
            raise
    
    def get_metadata(self, file_path: str) -> Dict[str, Any]:
        """Extract DOCX metadata."""
        try:
            doc = Document(file_path)
            return {
                'num_paragraphs': len(doc.paragraphs),
                'file_size': os.path.getsize(file_path),
                'format': 'DOCX'
            }
        except Exception as e:
            logger.error(f"Error extracting DOCX metadata: {e}")
            return {}


class TXTLoader(BaseDocumentLoader):
    """Loader for plain text documents."""
    
    def load(self, file_path: str) -> str:
        """Load text content from TXT file."""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read().strip()
        except Exception as e:
            logger.error(f"Error loading TXT {file_path}: {e}")
            raise
    
    def get_metadata(self, file_path: str) -> Dict[str, Any]:
        """Extract TXT metadata."""
        try:
            return {
                'file_size': os.path.getsize(file_path),
                'format': 'TXT'
            }
        except Exception as e:
            logger.error(f"Error extracting TXT metadata: {e}")
            return {}


class HTMLLoader(BaseDocumentLoader):
    """Loader for HTML documents."""
    
    def load(self, file_path: str) -> str:
        """Extract text content from HTML file."""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                soup = BeautifulSoup(file.read(), 'html.parser')
                return soup.get_text().strip()
        except Exception as e:
            logger.error(f"Error loading HTML {file_path}: {e}")
            raise
    
    def get_metadata(self, file_path: str) -> Dict[str, Any]:
        """Extract HTML metadata."""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                soup = BeautifulSoup(file.read(), 'html.parser')
                return {
                    'title': soup.title.string if soup.title else None,
                    'file_size': os.path.getsize(file_path),
                    'format': 'HTML'
                }
        except Exception as e:
            logger.error(f"Error extracting HTML metadata: {e}")
            return {}


class DocumentLoader:
    """Main document loader that handles multiple formats."""
    
    def __init__(self):
        """Initialize document loader with format handlers."""
        self.loaders = {
            '.pdf': PDFLoader(),
            '.docx': DOCXLoader(),
            '.txt': TXTLoader(),
            '.html': HTMLLoader(),
            '.htm': HTMLLoader()
        }
    
    def load_document(self, file_path: str) -> Dict[str, Any]:
        """Load document and return content with metadata."""
        file_path = Path(file_path)
        
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        file_extension = file_path.suffix.lower()
        
        if file_extension not in self.loaders:
            raise ValueError(f"Unsupported file format: {file_extension}")
        
        loader = self.loaders[file_extension]
        
        try:
            content = loader.load(str(file_path))
            metadata = loader.get_metadata(str(file_path))
            metadata['file_path'] = str(file_path)
            metadata['file_name'] = file_path.name
            
            return {
                'content': content,
                'metadata': metadata
            }
        except Exception as e:
            logger.error(f"Failed to load document {file_path}: {e}")
            raise
    
    def load_directory(self, directory_path: str) -> List[Dict[str, Any]]:
        """Load all supported documents from a directory."""
        directory = Path(directory_path)
        documents = []
        
        if not directory.exists():
            raise FileNotFoundError(f"Directory not found: {directory_path}")
        
        for file_path in directory.rglob('*'):
            if file_path.is_file() and file_path.suffix.lower() in self.loaders:
                try:
                    doc = self.load_document(str(file_path))
                    documents.append(doc)
                    logger.info(f"Loaded document: {file_path.name}")
                except Exception as e:
                    logger.warning(f"Failed to load {file_path}: {e}")
                    continue
        
        return documents
    
    def get_supported_formats(self) -> List[str]:
        """Get list of supported file formats."""
        return list(self.loaders.keys())
